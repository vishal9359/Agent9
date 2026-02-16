"""
Agent9 - C++ function flowchart generator (AST-driven, low-hallucination)

New approach (pseudo-code labels + fewer LLM calls):
- Control-flow is deterministic from libclang AST (if/loop/switch/try/return/break/continue).
- Process labels are pseudo-code close to the original statements (no "gist").
- LLM is OPTIONAL and, when enabled, is called in BATCHES to compress/pretty-print
  multiple process blocks at once (never one call per statement).

Mermaid constraints handled:
- Mermaid flowchart node labels must not contain raw bracket characters: () {} [].
  We escape them using Mermaid entity codes: #40; #41; #91; #93; #123; #125;.
- We also escape characters that often break parsing/rendering: ; : < >.
  (`<br/>` is preserved to support multi-line assignment blocks.)
"""

from __future__ import annotations

import argparse
import json
import os
import re
import subprocess
import time
from collections import defaultdict
from dataclasses import dataclass
from typing import Optional

from clang import cindex
from langchain.messages import HumanMessage
from langchain_ollama import ChatOllama
from docx import Document
from docx.shared import Inches


SUPPORTED_EXT = (".c", ".cpp", ".cc", ".cxx")
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DEFAULT_MERMAID_DIR = os.path.join(BASE_DIR, "mermaid_converter")
DEFAULT_OUT_DIR = os.path.join(BASE_DIR, "docs")

VERBOSE = True

# Mandatory libclang configuration (per requirements)
DEFAULT_LIBCLANG_PATH = "/usr/lib/llvm-18/lib/libclang.so"
try:
    cindex.Config.set_library_file(DEFAULT_LIBCLANG_PATH)
except Exception as e:  # pragma: no cover
    raise RuntimeError(
        f"Failed to configure libclang at {DEFAULT_LIBCLANG_PATH}. "
        "Install LLVM 18 libclang at that path (or adjust the code if your environment differs). "
        f"Original error: {e}"
    )


def log(msg: str):
    # Mandatory progress output (per requirements)
    print(msg, flush=True)


# Labeling limits
MAX_LABEL_CHARS = 180
MAX_BLOCK_CHARS_FOR_NODE = 1200  # stored for debugging only (LLM prompt uses BASE_LABEL)
LLM_CHUNK_CHAR_BUDGET = 6000     # keep batched prompts small for speed
LLM_SKIP_LABEL_MAX_CHARS = 90    # if base label is already short/clean, skip LLM rewrite

# Split long straight-line code into segments (avoid a single giant box)
PENDING_SEGMENT_MAX_LINES = 35
PENDING_SEGMENT_MAX_STMTS = 14

KEYWORD_TOKENS = {
    "if",
    "for",
    "while",
    "switch",
    "return",
    "new",
    "delete",
    "sizeof",
    "static_cast",
    "dynamic_cast",
    "const_cast",
    "reinterpret_cast",
}


def _ck(name: str):
    """CursorKind compatibility across clang bindings."""
    return getattr(cindex.CursorKind, name, None)


CK_CXX_TRY = _ck("CXX_TRY_STMT") or _ck("TRY_STMT")
CK_CXX_CATCH = _ck("CXX_CATCH_STMT")
CK_CXX_FOR_RANGE = _ck("CXX_FOR_RANGE_STMT")


def _kind_str(kind) -> str:
    return (getattr(kind, "name", None) or getattr(kind, "spelling", None) or str(kind) or "").upper()


def _is_range_for_cursor(cursor) -> bool:
    if cursor is None:
        return False
    k = _kind_str(getattr(cursor, "kind", ""))
    return ("FOR_RANGE" in k) or ("CXX_FOR_RANGE" in k)


def is_cpp_file(path: str) -> bool:
    return path.endswith(SUPPORTED_EXT)


def get_module_name(file_path: str, root_dir: str) -> str:
    rel = os.path.relpath(file_path, root_dir)
    no_ext = os.path.splitext(rel)[0]
    return ".".join(no_ext.split(os.sep))


def node_uid(cursor) -> str:
    loc = cursor.location
    return f"{cursor.spelling}:{loc.file.name}:{loc.line}"


def clean_unicode_chars(text: str) -> str:
    """Keep printable ASCII + newlines/tabs only."""
    if not text:
        return ""
    ascii_text = text.encode("ascii", "ignore").decode("ascii")
    ascii_text = re.sub(r"[^\x20-\x7E\n\r\t]", "", ascii_text)
    return ascii_text


def clamp_label(s: str) -> str:
    s = clean_unicode_chars(s or "")
    s = re.sub(r"\s+", " ", s).strip()
    if len(s) > MAX_LABEL_CHARS:
        s = s[: MAX_LABEL_CHARS - 3].rstrip()
        # Avoid leaving a dangling/incomplete Mermaid entity-code sequence like "#40"
        # (entity codes are plain text here, but truncating mid-sequence harms readability).
        s = re.sub(r"#\d{1,4}$", "", s).rstrip()
        s = s + "..."
    return s


def _escape_mermaid_chars(text: str) -> str:
    """
    Escape characters that commonly break Mermaid in node labels.
    Uses Mermaid entity-code syntax (not HTML entities): `#40;` etc.
    Preserves `<br/>` (multi-line label marker).
    """
    if not text:
        return text

    br_placeholder = "__MERMAID_BR__"
    text = text.replace("<br>", "<br/>").replace("<br/>", br_placeholder)

    # IMPORTANT:
    # - Only escape bracket characters that Mermaid uses as node delimiters: () {} [].
    # - Do NOT escape ';' or ':' here. Entity codes require ';', and ':' is needed for `std::`.
    esc = {
        "(": "#40;",
        ")": "#41;",
        "[": "#91;",
        "]": "#93;",
        "{": "#123;",
        "}": "#125;",
    }
    text = "".join(esc.get(ch, ch) for ch in text)
    return text.replace(br_placeholder, "<br/>")


def _strip_comments(code: str) -> str:
    """Strip // and /* */ comments (best-effort)."""
    if not code:
        return ""
    code = re.sub(r"/\*[\s\S]*?\*/", " ", code)
    code = re.sub(r"//.*?$", " ", code, flags=re.MULTILINE)
    return code


def _split_statements(code: str) -> list[str]:
    """
    Best-effort statement splitter for straight-line blocks.
    Splits on ';' at depth 0. Keeps tail fragment if non-empty.
    """
    code = clean_unicode_chars(code or "")
    code = _strip_comments(code).strip()
    if not code:
        return []

    out: list[str] = []
    depth = 0
    curr: list[str] = []
    in_str = False
    str_ch = ""

    for ch in code:
        if in_str:
            curr.append(ch)
            if ch == str_ch:
                in_str = False
            continue

        if ch in ("'", '"'):
            in_str = True
            str_ch = ch
            curr.append(ch)
            continue

        if ch in "([{":
            depth += 1
        elif ch in ")]}":
            depth = max(0, depth - 1)

        if ch == ";" and depth == 0:
            stmt = "".join(curr).strip()
            if stmt:
                out.append(stmt)
            curr = []
            continue

        curr.append(ch)

    tail = "".join(curr).strip()
    if tail:
        out.append(tail)
    return out


_DECL_ASSIGN_RE = re.compile(
    r"^\s*(?:const\s+)?(?:static\s+)?(?:inline\s+)?(?:volatile\s+)?"
    r"(?:unsigned\s+|signed\s+|long\s+|short\s+)?"
    r"[A-Za-z_][A-Za-z0-9_:<>]*\s+([A-Za-z_][A-Za-z0-9_]*)\s*=\s*(.+)$"
)


def _simplify_decl_assign(stmt: str) -> str:
    """Convert `Type name = expr` -> `name = expr` (best-effort)."""
    m = _DECL_ASSIGN_RE.match(stmt or "")
    if not m:
        return stmt
    return f"{m.group(1)} = {m.group(2).strip()}"


_DECL_CTOR_RE = re.compile(
    r"^\s*(?:const\s+)?(?:static\s+)?(?:inline\s+)?(?:volatile\s+)?"
    r"(?:unsigned\s+|signed\s+|long\s+|short\s+)?"
    r"[A-Za-z_][A-Za-z0-9_:<>]*\s+([A-Za-z_][A-Za-z0-9_]*)\s*\((.*)\)\s*$",
    flags=re.DOTALL,
)


def _simplify_decl_ctor(stmt: str) -> str:
    """
    Convert constructor-style declarations into pseudo-code:
      `Type var(args)` -> `var(args)`
    Drops the type (including templates) which is usually noise in flowcharts.
    """
    s = clean_unicode_chars(stmt or "").strip()
    m = _DECL_CTOR_RE.match(s)
    if not m:
        return stmt
    name = m.group(1)
    args = re.sub(r"\s+", " ", (m.group(2) or "").strip())
    return f"{name}({args})"


_PRESERVE_MEMBER_CONTEXT = {
    # Container/iterator calls where object context matters for clarity
    "size",
    "begin",
    "end",
    "erase",
    "push_back",
    "push_front",
    "insert",
    "find",
}


# Populated per parsed file; used to add purpose hints for call sites.
_GLOBAL_FN_PURPOSE: dict[str, str] = {}


def _split_identifier_words(name: str) -> list[str]:
    """
    Split a C/C++ identifier into human-ish words.
    Handles underscores and basic CamelCase.
    """
    name = clean_unicode_chars(name or "").strip()
    if not name:
        return []
    parts = re.split(r"[_\s]+", name)
    words: list[str] = []
    for p in parts:
        if not p:
            continue
        # Split camel case boundaries: FooBar -> Foo Bar, TRTRConfirm -> TRTR Confirm
        w = re.sub(r"([a-z0-9])([A-Z])", r"\1 \2", p)
        w = re.sub(r"([A-Z]+)([A-Z][a-z])", r"\1 \2", w)
        words.extend([x for x in w.split() if x])
    return words


_VERB_WORDS = {
    "get",
    "set",
    "update",
    "identify",
    "confirm",
    "enqueue",
    "dequeue",
    "handle",
    "process",
    "compute",
    "calculate",
    "build",
    "create",
    "init",
}


def _humanize_function_purpose(fn_name: str) -> str:
    """
    Derive a short purpose phrase from a function name.
    Example: TRTR_confirm_GetJCInfoTable -> Get JC Info Table
    """
    words = _split_identifier_words(fn_name)
    if not words:
        return ""
    low = [w.lower() for w in words]
    # Prefer "get/update/identify/..." even if earlier tokens like "confirm" exist.
    preferred = [
        "get",
        "update",
        "identify",
        "enqueue",
        "dequeue",
        "handle",
        "compute",
        "calculate",
        "confirm",
        "set",
        "init",
        "create",
        "build",
        "process",
    ]
    start_idx = 0
    for verb in preferred:
        if verb in low:
            start_idx = low.index(verb)
            break
    core = words[start_idx:]
    if not core:
        core = words
    # Title-ish casing but preserve all-caps tokens
    out = []
    for w in core:
        out.append(w if (w.isupper() and len(w) > 1) else w.capitalize())
    return " ".join(out).strip()


def _extract_first_comment_in_extent(file_lines: list[str], extent) -> str:
    """
    Extract the first inline comment inside a function extent, if present.
    """
    if not file_lines or not extent:
        return ""
    s = max(1, extent.start.line)
    e = min(len(file_lines), extent.end.line)
    # Search shortly after the signature
    for i in range(s, min(e, s + 12)):
        line = file_lines[i - 1].strip()
        if not line:
            continue
        if line.startswith("//"):
            txt = line[2:].strip()
            return txt
        if "/*" in line:
            # best-effort: strip /* ... */
            txt = re.sub(r"/\*+|\*+/|\*/", " ", line)
            txt = re.sub(r"\s+", " ", txt).strip()
            return txt
    return ""


def _normalize_member_calls(text: str) -> str:
    """
    Normalize member-function calls into a Mermaid-friendly pseudo form:

      obj->Method(args)  ->  obj.<Method>(args)
      obj.Method(args)   ->  obj.<Method>(args)

    This avoids putting raw '->' into labels and keeps the exact parameter text.
    We do NOT attempt to parse/modify nested calls or parameter expressions.
    """
    if not text:
        return text

    def repl(match: re.Match) -> str:
        obj = match.group(1)
        method = match.group(2)
        # Preserve object context for common getter/predicate patterns so conditions
        # can be rewritten into readable English (e.g., ubio.IsSyncMode()).
        if method in _PRESERVE_MEMBER_CONTEXT or re.match(r"^(Is|Get|Set|Has|Can|Do)[A-Z_]", method):
            return f"{obj}.{method}("
        return f"{method}("

    # Only rewrite when it's clearly a call (followed by '(').
    return re.sub(
        r"\b([A-Za-z_][A-Za-z0-9_:]*)\s*(?:->|\.)\s*([A-Za-z_][A-Za-z0-9_]*)\s*\(",
        repl,
        text,
    )


def _apply_nbsp_around_assignment_ops(label: str) -> str:
    """
    In multi-line labels, prevent ugly wrap around '=' using non-breaking spaces.
    """
    # Disabled by design: requirements specify using entity codes only for brackets,
    # not for whitespace. Keep plain ASCII spaces around '='.
    return label


def clean_label_text(label: str, for_condition: bool = False) -> str:
    """
    Convert raw code/pseudo-code into Mermaid-safe label text.
    """
    label = clean_unicode_chars(label or "")
    label = label.replace("<br>", "<br/>")
    br = "__BR__"
    label = label.replace("<br/>", br)

    # Keep labels close to the original pseudo-code (no "gist").
    # Only normalize member calls for readability; do not drop/alter parameters.
    label = _normalize_member_calls(label)
    label = label.replace("&&", " and ").replace("||", " or ")
    label = _replace_comparisons_with_words(label)

    # Normalize whitespace but preserve <br/>
    label = re.sub(r"\s+", " ", label).strip()
    label = label.replace(br, "<br/>")

    label = _escape_mermaid_chars(label)

    return clamp_label(label) or "Process block"


_CMP_OP_WORDS = {
    "==": "equals",
    "!=": "not equal",
    "<=": "less than or equal to",
    ">=": "greater than or equal to",
    "<": "less than",
    ">": "greater than",
}


def _replace_comparisons_with_words(expr: str) -> str:
    """
    Replace comparison operators with words in a best-effort manner.
    Also normalizes `A > B` into `B is less than A` (no '>' operator).
    """
    s = clean_unicode_chars(expr or "")
    if not s:
        return ""

    # Work on a copy; avoid touching things like templates by requiring word-ish sides.
    pat = re.compile(
        r"(\b[A-Za-z_][A-Za-z0-9_:\.\[\]]*(?:\s*\([^)]*\))?)\s*(==|!=|>=|<=|>|<)\s*"
        r"(\b[A-Za-z_][A-Za-z0-9_:\.\[\]]*(?:\s*\([^)]*\))?)"
    )

    def repl(m: re.Match) -> str:
        left = m.group(1).strip()
        op = m.group(2)
        right = m.group(3).strip()
        def looks_like_constant(x: str) -> bool:
            x2 = re.sub(r"\s+", "", x)
            return (
                bool(re.fullmatch(r"[A-Z0-9_:\.]+", x2))
                or bool(re.fullmatch(r"\d+", x2))
                or x2.lower() in ("true", "false", "nullptr", "null")
            )

        # Prefer keeping the variable-ish side first when comparing against constants
        swap = looks_like_constant(left) and not looks_like_constant(right)
        if swap and op in ("==", "!="):
            left, right = right, left

        rlow = right.lower().strip()
        # Special-case unknown enum comparisons: event == BackendEvent_Unknown -> event is unknown
        if op in ("==", "!=") and re.search(r"(?:^|_)unknown$", right, flags=re.IGNORECASE):
            return f"{left} is {'not ' if op == '!=' else ''}unknown"

        # Special-case null checks
        if op in ("==", "!=") and rlow in ("nullptr", "null"):
            return f"{left} is {'not ' if op == '!=' else ''}null"

        # Special-case boolean literal comparisons
        if op in ("==", "!=") and rlow in ("true", "false"):
            want_true = (rlow == "true")
            positive = (op == "==" and want_true) or (op == "!=" and not want_true)
            return f"{left} is {'true' if positive else 'false'}"

        if op == ">":
            return f"{right} is {_CMP_OP_WORDS['<']} {left}" if swap else f"{left} is {_CMP_OP_WORDS['>']} {right}"
        if op == ">=":
            return f"{right} is {_CMP_OP_WORDS['<=']} {left}" if swap else f"{left} is {_CMP_OP_WORDS['>=']} {right}"
        if op == "<":
            return f"{left} is {_CMP_OP_WORDS['<']} {right}"
        if op == "<=":
            return f"{left} is {_CMP_OP_WORDS['<=']} {right}"
        if op == "==":
            return f"{left} {_CMP_OP_WORDS['==']} {right}"
        if op == "!=":
            return f"{left} {_CMP_OP_WORDS['!=']} {right}"
        return m.group(0)

    prev = None
    while prev != s:
        prev = s
        s = pat.sub(repl, s)
    return s


def clean_condition_text(label: str) -> str:
    """
    Clean condition expressions for decision nodes.
    Keep the exact condition; remove leading keyword and removable outer parentheses.
    """
    label = clean_unicode_chars(label or "")
    label = re.sub(r"^\s*(if|for|while|switch)\b", "", label, flags=re.IGNORECASE).strip()
    br = "__BR__"
    label = label.replace("<br/>", br)

    label = label.strip()
    if label.startswith("(") and label.endswith(")"):
        depth = 0
        is_outer = True
        for i, ch in enumerate(label):
            if ch == "(":
                depth += 1
            elif ch == ")":
                depth -= 1
                if depth == 0 and i < len(label) - 1:
                    is_outer = False
                    break
        if is_outer:
            label = label[1:-1].strip()

    # Normalize member call display and boolean ops
    label = _normalize_member_calls(label)
    label = label.replace("&&", " and ").replace("||", " or ")

    def _predicate_phrase(method: str) -> str:
        # Generic: IsReady -> ready, IsValid -> valid, IsSomethingMode -> in something mode, IsFlagSet -> flag is set
        core = re.sub(r"^Is", "", method).strip()
        if not core:
            return method
        words = _split_identifier_words(core)
        low = [w.lower() for w in words]
        if low and low[-1] == "mode":
            head = " ".join(words[:-1]).strip() or core
            return f"in {head} mode"
        if low and low[-1] == "set":
            head = " ".join(words[:-1]).strip() or core
            return f"{head} is set"
        return " ".join(words).strip() or core

    def _rewrite_predicate_comparisons(s: str) -> str:
        out = s
        # obj.IsX() == true/false  ->  obj is (not) <phrase>
        def repl_pred(m: re.Match) -> str:
            obj = m.group(1)
            method = m.group(2)
            op = m.group(3)
            lit = (m.group(4) or "").lower()
            want_true = (lit == "true")
            positive = (op == "==" and want_true) or (op == "!=" and not want_true)
            phr = _predicate_phrase(method)
            if phr.lower().endswith(" is set"):
                return f"{obj} {phr}" if positive else f"{obj} {phr.replace(' is set', ' is not set')}"
            return f"{obj} is {phr}" if positive else f"{obj} is not {phr}"

        out = re.sub(
            r"\b([A-Za-z_][A-Za-z0-9_]*)\.(Is[A-Za-z_][A-Za-z0-9_]*)\s*\([^)]*\)\s*(==|!=)\s*(true|false)\b",
            repl_pred,
            out,
            flags=re.IGNORECASE,
        )

        # !Func(...) -> Func(...) == false (best-effort; keeps semantics in label text)
        out = re.sub(
            r"!\s*([A-Za-z_][A-Za-z0-9_:]*\s*\([^)]*\))",
            r"\1 == false",
            out,
        )
        return out

    label = _rewrite_predicate_comparisons(label)

    # Remove operator symbols (==, !=, <, >, <=, >=) in favor of words
    label = _replace_comparisons_with_words(label)

    # Best-effort natural language for common zero/pending patterns
    label = re.sub(
        r"\b([A-Za-z_][A-Za-z0-9_]*(?:\[[^\]]+\])*)\s+equals\s+0\b",
        lambda m: "no " + _humanize_var(re.sub(r"\[[^\]]+\]", "", m.group(1))) + " exists",
        label,
        flags=re.IGNORECASE,
    )

    # (No hardcoded predicate helpers here; keep logic generic)

    # Whitespace normalize but preserve <br/>
    label = re.sub(r"\s+", " ", label).strip()
    label = label.replace(br, "<br/>")

    label = _escape_mermaid_chars(label)
    return clamp_label(label) or "condition"


def extract_source_lines(file_lines: list[str], start_line: int, end_line: int) -> list[str]:
    start_line = max(1, start_line)
    end_line = max(start_line, end_line)
    return [l.rstrip("\n") for l in file_lines[start_line - 1 : min(len(file_lines), end_line)]]


def cursor_extent_text(cursor, file_lines: list[str]) -> str:
    """
    Extract exact source text for a cursor extent (respects start/end columns).
    This prevents capturing trailing braces on the same line for conditions.
    """
    if not cursor or not getattr(cursor, "extent", None):
        return ""
    try:
        sline = max(1, cursor.extent.start.line)
        eline = max(1, cursor.extent.end.line)
        scol = max(1, cursor.extent.start.column)
        ecol = max(1, cursor.extent.end.column)

        if not file_lines:
            return ""
        if sline > len(file_lines) or eline > len(file_lines):
            return ""

        if sline == eline:
            line = file_lines[sline - 1].rstrip("\n")
            return line[scol - 1 : max(scol - 1, ecol - 1)].strip()

        parts: list[str] = []
        first = file_lines[sline - 1].rstrip("\n")
        parts.append(first[scol - 1 :])
        for i in range(sline, eline - 1):
            parts.append(file_lines[i].rstrip("\n"))
        last = file_lines[eline - 1].rstrip("\n")
        parts.append(last[: max(0, ecol - 1)])
        return "\n".join(parts).strip()
    except Exception:
        return ""


def cursor_text(cursor, file_lines: list[str]) -> str:
    if not cursor or not getattr(cursor, "extent", None):
        return ""
    return "\n".join(extract_source_lines(file_lines, cursor.extent.start.line, cursor.extent.end.line)).strip()


def _function_def_snippet(fn_cursor, file_lines: list[str]) -> str:
    """
    Build a compact, LLM-friendly snippet for a function definition.
    Used to infer the callee's purpose without generating its internal flowchart.
    """
    try:
        if not fn_cursor or not getattr(fn_cursor, "extent", None):
            return ""
        name = getattr(fn_cursor, "spelling", "") or ""
        if not name:
            return ""
        start = max(1, fn_cursor.extent.start.line)
        end = min(len(file_lines), fn_cursor.extent.end.line)
        lines = extract_source_lines(file_lines, start, end)
        if not lines:
            return ""

        # Keep signature + a small window + first return statement
        head = [l.rstrip() for l in lines[:50] if l.strip()]
        snippet = "\n".join(head)

        # Add the first return line if it appears later
        if "return" not in snippet:
            for l in lines:
                if re.search(r"\breturn\b", l):
                    snippet = snippet + "\n" + l.strip()
                    break

        return clean_unicode_chars(snippet).strip()
    except Exception:
        return ""


def _extract_paren_group_after(keyword: str, text: str) -> str:
    """
    Extract the parenthesized group that follows a keyword, handling nested ().
    Example: for (a=begin(); b!=end(); ++b) -> returns "a=begin(); b!=end(); ++b"
    """
    if not text:
        return ""
    m = re.search(rf"\b{re.escape(keyword)}\b", text)
    if not m:
        return ""
    i = m.end()
    n = len(text)
    while i < n and text[i].isspace():
        i += 1
    if i >= n or text[i] != "(":
        return ""
    i += 1
    depth = 1
    start = i
    in_str = False
    str_ch = ""
    while i < n:
        ch = text[i]
        if in_str:
            if ch == str_ch:
                in_str = False
            i += 1
            continue
        if ch in ("'", '"'):
            in_str = True
            str_ch = ch
            i += 1
            continue
        if ch == "(":
            depth += 1
        elif ch == ")":
            depth -= 1
            if depth == 0:
                return text[start:i]
        i += 1
    return ""


def _split_top_level_semicolons(s: str) -> list[str]:
    """Split a string by ';' at depth 0 (handles nested ()/[]/{})."""
    if not s:
        return []
    out: list[str] = []
    depth = 0
    curr: list[str] = []
    in_str = False
    str_ch = ""
    for ch in s:
        if in_str:
            curr.append(ch)
            if ch == str_ch:
                in_str = False
            continue
        if ch in ("'", '"'):
            in_str = True
            str_ch = ch
            curr.append(ch)
            continue
        if ch in "([{":
            depth += 1
        elif ch in ")]}":
            depth = max(0, depth - 1)
        if ch == ";" and depth == 0:
            part = "".join(curr).strip()
            if part:
                out.append(part)
            curr = []
            continue
        curr.append(ch)
    tail = "".join(curr).strip()
    if tail:
        out.append(tail)
    return out


def _describe_member_chain(expr: str) -> str:
    """
    Convert simple member/index chains into readable text.
    Example: jcInfoTable->astInfoTable[ntr].next -> next of astInfoTable at index ntr of jcInfoTable
    Best-effort; used mainly for loop bounds and pseudo labels.
    """
    s = clean_unicode_chars(expr or "").strip()
    if not s:
        return ""
    s = s.replace("->", ".")
    parts = [p for p in s.split(".") if p]
    if not parts:
        return s

    def fmt_part(p: str) -> str:
        m = re.match(r"^([A-Za-z_][A-Za-z0-9_]*)\s*\[\s*([^\]]+)\s*\]\s*$", p)
        if m:
            return f"{m.group(1)} at index {m.group(2).strip()}"
        return p.strip()

    chain = [fmt_part(p) for p in parts]
    # Build "last of prev of base" wording
    out = chain[-1]
    for prev in reversed(chain[:-1]):
        out = f"{out} of {prev}"
    return out


def _humanize_var(name: str) -> str:
    words = _split_identifier_words(name)
    if not words:
        return name
    # Lowercase first token if it looks like a prefix (e.g. trType -> tr type)
    out = []
    for w in words:
        out.append(w.lower() if (len(w) <= 3 and w.isalpha()) else w)
    return " ".join(out)

_LAMBDA_DEF_RE = re.compile(
    r"(?s)\b([A-Za-z_][A-Za-z0-9_]*)\s*=\s*\[[^\]]*\]\s*(?:\([^)]*\))?\s*\{"
)


def _looks_like_lambda_definition(stmt: str) -> bool:
    if not stmt:
        return False
    return bool(re.search(r"(?s)\=\s*\[[^\]]*\]\s*(?:\([^)]*\))?\s*\{", stmt))


def _compress_lambda_definition(stmt: str) -> str:
    """
    Avoid dumping the full lambda body into a flowchart label.
    Convert `auto helper = [&](int x){ return x*2; };` -> `helper = lambda`.
    """
    s = clean_unicode_chars(stmt or "").strip()
    if not _looks_like_lambda_definition(s):
        return s
    m = _LAMBDA_DEF_RE.search(s)
    if m:
        name = m.group(1)
        return f"{name} = lambda"
    return "lambda"


_NOISE_STMT_RE = re.compile(
    r"^\s*(?:std::)?(?:cout|cerr|clog)\s*<<|"
    r"\b(?:printf|fprintf|sprintf|snprintf|puts|putchar|getchar|perror)\s*\(",
    flags=re.IGNORECASE,
)


def _is_noise_statement(stmt: str) -> bool:
    """
    Statements that usually add noise to control-flow flowcharts.
    Keep this conservative: only obvious I/O.
    """
    if not stmt:
        return True
    return bool(_NOISE_STMT_RE.search(stmt))


def _extract_call_expressions(stmt: str) -> list[str]:
    """
    Extract call expressions like `foo(x, y)` from a statement.
    Best-effort, not a full C++ parser.
    """
    s = clean_unicode_chars(stmt or "")
    out: list[str] = []

    i = 0
    n = len(s)
    in_str = False
    str_ch = ""

    def is_ident_char(ch: str) -> bool:
        return ch.isalnum() or ch == "_" or ch == ":"

    while i < n:
        ch = s[i]
        if in_str:
            if ch == str_ch:
                in_str = False
            i += 1
            continue
        if ch in ("'", '"'):
            in_str = True
            str_ch = ch
            i += 1
            continue

        if (ch.isalpha() or ch == "_") and (i == 0 or not is_ident_char(s[i - 1])):
            j = i + 1
            while j < n and is_ident_char(s[j]):
                j += 1
            name = s[i:j]
            k = j
            while k < n and s[k].isspace():
                k += 1
            if k < n and s[k] == "(" and name not in KEYWORD_TOKENS:
                depth = 0
                t = k
                in_str2 = False
                str_ch2 = ""
                while t < n:
                    c2 = s[t]
                    if in_str2:
                        if c2 == str_ch2:
                            in_str2 = False
                        t += 1
                        continue
                    if c2 in ("'", '"'):
                        in_str2 = True
                        str_ch2 = c2
                        t += 1
                        continue
                    if c2 == "(":
                        depth += 1
                    elif c2 == ")":
                        depth -= 1
                        if depth == 0:
                            out.append(s[i : t + 1].strip())
                            i = t + 1
                            break
                    t += 1
                else:
                    i = j
                    continue
                continue
        i += 1

    seen = set()
    uniq: list[str] = []
    for c in out:
        if c not in seen:
            uniq.append(c)
            seen.add(c)
    return uniq


_ASSIGN_RE = re.compile(
    r"^\s*([A-Za-z_][A-Za-z0-9_:\.\[\]]*)\s*=\s*(.+)$"
)


def _simplify_pseudocode_statement(stmt: str) -> str:
    """
    Apply pseudo-code rules to a single statement (best-effort):
    - If it's an assignment with a call on RHS, prefer showing the RHS call only.
    - Keep increments, indexing, arithmetic, and plain assignments as-is.
    """
    s = clean_unicode_chars(stmt or "").strip()
    if not s:
        return ""

    # Do not touch comparisons or compound operators here.
    if "==" in s or ">=" in s or "<=" in s or "!=" in s:
        return s

    m = _ASSIGN_RE.match(s)
    if m:
        lhs = m.group(1).strip()
        rhs = m.group(2).strip()
        rhs = _normalize_member_calls(rhs)
        # If RHS is a member-context call like container.begin()/end()/size(),
        # keep the full assignment for clarity (pseudo-code).
        if re.search(r"\b[A-Za-z_][A-Za-z0-9_]*\.(?:begin|end|size)\s*\(", rhs):
            return s
        calls = _extract_call_expressions(rhs)
        if calls:
            # Preserve assignment semantics for downstream LLM labeling.
            return f"{lhs} = {calls[0]}"
        return s

    return s


def build_pseudocode_label(block_text: str) -> str:
    """
    Deterministic pseudo-code label builder (no LLM).
    """
    block_text = clean_unicode_chars(block_text or "")
    if not block_text.strip():
        return "Process block"

    stmts: list[str] = []
    had_lambda_def = False
    for s in _split_statements(block_text):
        s2 = s.strip()
        if not s2:
            continue
        if _looks_like_lambda_definition(s2):
            had_lambda_def = True
            # Keep a short marker only if this is the only thing in the block.
            s2 = _compress_lambda_definition(s2)
        s2 = _simplify_decl_ctor(s2)
        s2 = _simplify_decl_assign(s2)
        s2 = _simplify_pseudocode_statement(s2)
        s2 = s2.strip()
        if s2:
            stmts.append(s2)

    if not stmts:
        return clean_label_text(re.sub(r"\s+", " ", block_text).strip()) or "Process block"

    # Filter obvious noise (e.g. cout << ...)
    filtered = [s for s in stmts if not _is_noise_statement(s)]
    if not filtered:
        filtered = stmts

    # If there are meaningful calls in this block, prefer showing calls only.
    calls: list[str] = []
    for s in filtered:
        calls.extend(_extract_call_expressions(_normalize_member_calls(s)))

    # Special case: blocks that define a lambda and then call something.
    # Drop the lambda definition from the label and show the call(s).
    if had_lambda_def and calls:
        label = "<br/>".join(calls[:4])
        if len(calls) > 4:
            label += "<br/>..."
    else:
        simplified: list[str] = []
        for s in filtered[:6]:
            s2 = _simplify_pseudocode_statement(s).strip()
            if s2:
                simplified.append(s2)
        label = "<br/>".join(simplified)
        if len(filtered) > 6:
            label += "<br/>..."
    return clean_label_text(label)


class BatchLLMLabeler:
    """
    Optional batched LLM labeler that compresses/pretty-prints pseudo-code.
    It should not invent logic.
    """

    def __init__(self, llm, enabled: bool):
        self.llm = llm
        self.enabled = bool(enabled and llm and HumanMessage)
        self.cache: dict[str, str] = {}
        # LLM-inferred purposes for called functions (by function name)
        self.purpose_cache: dict[str, str] = {}

    def _ensure_purposes(self, fn_names: list[str]) -> None:
        """
        Infer purpose phrases for functions using their definitions (snippets).
        This lets us "look inside" a callee for understanding, but we do not
        generate a flowchart for the callee's internals.
        """
        if not self.enabled:
            return

        todo = []
        for nm in fn_names:
            if not nm or nm in self.purpose_cache:
                continue
            snip = (_GLOBAL_FN_DEF_SNIPPET.get(nm) or "").strip()
            if snip:
                todo.append(nm)
        if not todo:
            return

        chunk: list[str] = []
        chars = 0

        def flush():
            nonlocal chunk, chars
            if not chunk:
                return

            blocks = []
            for nm in chunk:
                snip = (_GLOBAL_FN_DEF_SNIPPET.get(nm) or "").strip()
                snip_lines = [l.rstrip() for l in snip.splitlines() if l.strip()]
                snip_short = "\n".join(snip_lines[:50])
                blocks.append(f"### FUNCTION {nm}\n{snip_short}\n")

            prompt = (
                "You summarize C/C++ functions into short purpose phrases for flowchart labels.\n"
                "Return STRICT JSON object mapping function names to purpose phrases.\n"
                "Rules:\n"
                "- 4-14 words\n"
                "- prefer verbs like Get/Update/Identify/Confirm/Enqueue/Dequeue/Handle/Compute\n"
                "- do NOT use operator symbols: == != > < <= >=\n"
                "- do NOT invent details not supported by the snippet\n"
                "\n"
                f"{''.join(blocks)}\n"
                "JSON:"
            )
            log("[LLM] Purpose prompt:\n" + prompt)
            t0 = time.perf_counter()
            resp = self.llm.invoke([HumanMessage(prompt)])
            dt = time.perf_counter() - t0
            log(f"[LLM] Purpose completed in {dt:.3f}s")
            try:
                data = json.loads((resp.content or "").strip())
                if isinstance(data, dict):
                    for k, v in data.items():
                        kk = str(k).strip()
                        vv = clean_unicode_chars(str(v)).strip()
                        if kk and vv:
                            self.purpose_cache[kk] = vv
            except Exception:
                pass

            chunk = []
            chars = 0

        for nm in todo:
            if chunk and (chars + len(nm) > 6500):
                flush()
            chunk.append(nm)
            chars += len(nm)
        flush()

    def label_many(self, blocks: list[tuple[str, str, str]]) -> dict[str, str]:
        out: dict[str, str] = {}
        if not blocks:
            return out

        def needs_llm(base_label: str) -> bool:
            """
            LLM is only used to lightly rewrite labels for readability.
            If the deterministic pseudo label is already short and clean,
            we can skip the rewrite with no impact on correctness.
            """
            bl = clean_unicode_chars(base_label or "").strip()
            if not bl:
                return True
            # Force LLM rewrite for anything that still looks like raw pseudo-code.
            if "<br/>" in bl:
                return True
            if re.search(r"(;|\+\+|--|\breturn\b|(?<![=!<>])=(?![=])|\(|\[|\]|->|\.)", bl):
                return True

            # Skip LLM only for already-English short labels.
            if len(bl) <= LLM_SKIP_LABEL_MAX_CHARS and "..." not in bl and re.match(
                r"^(Get|Update|Initialize|Identify|Fetch|Store|Increase|Decrease|Submit|Notify|Enqueue|Dequeue|Assign|Check|Iterate|Exit)\b",
                bl,
                flags=re.IGNORECASE,
            ):
                return False
            if bl.lower() in ("process block", "no operation"):
                return True
            return True

        remaining: list[tuple[str, str, str]] = []
        for bid, base_label, txt in blocks:
            # Cache by base label (code can vary but pseudo label might repeat)
            key = re.sub(r"\s+", " ", clean_unicode_chars(base_label or "")).strip()
            if key in self.cache:
                out[bid] = self.cache[key]
            else:
                remaining.append((bid, base_label, txt))

        if not remaining or not self.enabled:
            for bid, base_label, _txt in remaining:
                out[bid] = clean_label_text(base_label)
            return out

        # Use LLM to infer purposes for callees referenced in these labels.
        called: set[str] = set()
        for _bid, base_label, _txt in remaining:
            for call in _extract_call_expressions(base_label or ""):
                m = re.match(r"^([A-Za-z_][A-Za-z0-9_:]*)\s*\(", call)
                if m:
                    called.add(m.group(1))
        self._ensure_purposes(sorted(called))

        # Chunk to keep prompts sane
        chunk: list[tuple[str, str, str]] = []
        chunk_chars = 0

        def flush_chunk():
            nonlocal chunk, chunk_chars
            if not chunk:
                return
            mapping = self._invoke_chunk(chunk)
            for bid, base_label, txt in chunk:
                key = re.sub(r"\s+", " ", clean_unicode_chars(base_label or "")).strip()
                # If LLM fails for any block, keep deterministic base label.
                proposed = mapping.get(bid) or base_label
                # Guardrail: do not allow the LLM to drop statement lines.
                if base_label.count("<br/>") and proposed.count("<br/>") < base_label.count("<br/>"):
                    lbl = base_label
                else:
                    lbl = proposed
                lbl = clean_label_text(lbl)
                out[bid] = lbl
                self.cache[key] = lbl
            chunk = []
            chunk_chars = 0

        for bid, base_label, txt in remaining:
            # Skip LLM for already-good labels (speed)
            if not needs_llm(base_label):
                lbl = clean_label_text(base_label)
                out[bid] = lbl
                key = re.sub(r"\s+", " ", clean_unicode_chars(base_label or "")).strip()
                self.cache[key] = lbl
                continue

            # Keep code only for debugging/troubleshooting; do not include in prompt.
            t = clean_unicode_chars(txt or "")[:MAX_BLOCK_CHARS_FOR_NODE]
            if chunk and (chunk_chars + len(base_label) > LLM_CHUNK_CHAR_BUDGET):
                flush_chunk()
            chunk.append((bid, base_label, t))
            chunk_chars += len(base_label)

        flush_chunk()
        return out

    def _invoke_chunk(self, chunk: list[tuple[str, str, str]]) -> dict[str, str]:
        blocks_txt = []

        def infer_statement_type(base_label: str) -> str:
            bl = clean_unicode_chars(base_label or "").strip()
            if not bl:
                return "FUNCTION_CALL"
            # Prefer the first line/statement for classification.
            first = bl.split("<br/>", 1)[0].strip()
            s = first or bl
            s_l = s.lower()

            if re.search(r"\b(for each|after loop|loop|while|for)\b", s_l):
                return "LOOP"
            if s_l.startswith("check ") or re.search(r"(==|!=|>=|<=|<|>)", s):
                return "CONDITION"
            if re.search(r"(?<![=!<>])=(?![=])", s):
                return "ASSIGNMENT"
            if re.search(r"\b[A-Za-z_][A-Za-z0-9_:]*\s*\(", s):
                return "FUNCTION_CALL"
            return "FUNCTION_CALL"

        def _suggest_label_for_base_label(base_label: str, purpose_map: dict[str, str]) -> str:
            """
            Best-effort deterministic English suggestion to stabilize LLM output.
            Keeps one statement per line (uses literal '<br/>' separators).
            """
            bl = clean_unicode_chars(base_label or "").strip()
            if not bl:
                return ""

            def first_call_name(text: str) -> str:
                m = re.search(r"\b([A-Za-z_][A-Za-z0-9_:]*)\s*\(", text or "")
                return m.group(1) if m else ""

            def suggest_one(stmt: str) -> str:
                s = clean_unicode_chars(stmt or "").strip().rstrip(";")
                if not s:
                    return ""

                if s == "return" or s.startswith("return "):
                    return "Exit function"

                # ++ / --
                if s.endswith("++"):
                    v = s[:-2].strip()
                    v = re.sub(r"\[[^\]]+\]", "", v)
                    return f"Increment {_humanize_var(v)}"
                if s.endswith("--"):
                    v = s[:-2].strip()
                    v = re.sub(r"\[[^\]]+\]", "", v)
                    return f"Decrease {_humanize_var(v)}"

                m = _ASSIGN_RE.match(s)
                if m:
                    lhs = m.group(1).strip()
                    rhs = m.group(2).strip()
                    rhs = _normalize_member_calls(rhs)
                    lhs_base = re.sub(r"\[[^\]]+\]", "", lhs).strip()
                    lhs_desc = _humanize_var(lhs_base)
                    # If LHS is a member/index target, describe it more explicitly.
                    lhs_target = lhs_desc
                    m_lhs = re.match(r"^([A-Za-z_][A-Za-z0-9_]*)\s*(->|\.)\s*([A-Za-z_][A-Za-z0-9_]*)", lhs)
                    if m_lhs:
                        obj = m_lhs.group(1)
                        acc = m_lhs.group(2)
                        field = m_lhs.group(3)
                        idxs = re.findall(r"\[([^\]]+)\]", lhs)
                        lhs_target = f"{field} of {'pointer structure ' if acc == '->' else ''}{obj}".strip()
                        if idxs:
                            lhs_target += f" at index {idxs[0].strip()}"

                    # init to 0/null
                    if re.fullmatch(r"(?:0|0U|0u|NULL|nullptr)", rhs):
                        if "null" in rhs.lower():
                            return f"Initialize {lhs_target} to null"
                        return f"Initialize {lhs_target} to zero"

                    # x = x + y
                    if re.match(rf"^{re.escape(lhs_base)}\s*\+\s*", rhs):
                        y = re.sub(rf"^{re.escape(lhs_base)}\s*\+\s*", "", rhs).strip()
                        m_sz = re.search(r"\b([A-Za-z_][A-Za-z0-9_]*)\.GetSize\s*\([^)]*\)\s*$", y)
                        if m_sz:
                            return f"Increase {_humanize_var(lhs_base)} by {m_sz.group(1)} size"
                        if re.search(r"\bGetSize\s*\(", y):
                            return f"Increase {_humanize_var(lhs_base)} by size"
                        return f"Increase {_humanize_var(lhs_base)} by {y}"

                    # assignment from call
                    fn = first_call_name(rhs)
                    if fn:
                        pur = (purpose_map.get(fn) or _humanize_function_purpose(fn) or "").strip()
                        if pur:
                            return f"{pur} by calling {rhs}"
                        return f"Update {lhs_target} by calling {rhs}"

                    # assignment from member/index chain
                    if ("." in rhs) or ("[" in rhs):
                        desc = _describe_member_chain(rhs) or rhs
                        return f"Update {lhs_target} with {desc}"

                    return f"Update {lhs_target}"

                # plain call
                fn = first_call_name(s)
                if fn:
                    pur = (purpose_map.get(fn) or _humanize_function_purpose(fn) or "").strip()
                    return pur or f"Call {fn}"

                return s

            lines = [x.strip() for x in bl.split("<br/>") if x.strip()]
            out_lines: list[str] = []
            for x in lines:
                sug = (suggest_one(x) or "").strip()
                if not sug:
                    # Preserve line count; fall back to a compact cleaned statement.
                    sug = re.sub(r"\s+", " ", clean_unicode_chars(x)).strip() or "Step"
                out_lines.append(sug)
            return "<br/>".join(out_lines)

        for bid, base_label, txt in chunk:
            calls = _extract_call_expressions(base_label or "")
            ctx_lines = []
            purpose_map: dict[str, str] = {}
            for c in calls[:6]:
                m = re.match(r"^([A-Za-z_][A-Za-z0-9_:]*)\s*\(", c)
                if not m:
                    continue
                fn = m.group(1)
                pur = (self.purpose_cache.get(fn) or "").strip()
                if pur:
                    ctx_lines.append(f"- {fn}: {pur}")
                    purpose_map[fn] = pur
            ctx = "\n".join(ctx_lines) if ctx_lines else "none"
            stype = infer_statement_type(base_label)
            suggested = _suggest_label_for_base_label(base_label, purpose_map) or "none"
            blocks_txt.append(
                f"### BLOCK {bid}\n"
                f"STATEMENT_TYPE: {stype}\n"
                f"BASE_LABEL:\n{base_label}\n"
                f"CALLEE_PURPOSES:\n{ctx}\n"
                f"SUGGESTED_LABEL:\n{suggested}\n"
            )

        prompt = (
            """
Convert the following C/C++ code statement into a short, flowchart-friendly English label.

You are rewriting AST-derived pseudo-code, NOT raw source code.

STRICT RULES:
- Output must be suitable for a flowchart node
- Do NOT invent domain-specific or business meaning
- Light semantic inference from function and variable names is ALLOWED
- Preserve assignment direction and data flow
- Use simple verbs only: Get, Update, Initialize, Identify, Fetch, Store, Increase, Decrease, Submit, Notify, Enqueue, Dequeue, Assign, Check, Iterate, Exit
- Mention function calls explicitly using "by calling <function>()"
- Mention structure, pointer, and index access explicitly
- Each line should be one short sentence (multi-line uses literal '<br/>')
- Keep each line <= 140 characters

STATEMENT TYPE RULES:
- ASSIGNMENT:
  - Mention destination first
  - Describe source after "with" or "by calling"
- FUNCTION_CALL:
  - Describe action using function name
- LOOP:
  - Use quantifier form (For all / For each)
- CONDITION:
  - Use "Check ..."

OUTPUT FORMAT:
- Return STRICT JSON object mapping block ids to labels
- Preserve "<br/>" line count exactly
- Do NOT drop any statement line

IMPORTANT:
- If SUGGESTED_LABEL is not "none" and matches BASE_LABEL semantics, output it exactly.
- If CALLEE_PURPOSES provides a purpose for a called function, prefer that phrasing.

Do NOT use operator symbols: == != > < <= >=
Use words instead.

Code blocks:
"""
            f"{''.join(blocks_txt)}\n"
            "JSON:"
        )
        try:
            log("[LLM] Prompt:\n" + prompt)
            t0 = time.perf_counter()
            resp = self.llm.invoke([HumanMessage(prompt)])
            dt = time.perf_counter() - t0
            log(f"[LLM] Completed in {dt:.3f}s")
            content = (resp.content or "").strip()
            data = json.loads(content)
            if isinstance(data, dict):
                return {str(k): str(v) for k, v in data.items()}
        except Exception:
            return {}
        return {}


@dataclass
class Graph:
    nodes: dict[str, dict]  # id -> {"shape": str, "label": str}
    edges: list[tuple[str, str, str]]  # (src, dst, label)


class FlowBuilder:
    """
    Deterministic flow builder from AST statements.
    Basic blocks = grouped consecutive non-control statements.
    """

    def __init__(self, file_lines: list[str], root_dir: str, file_path: str, labeler: BatchLLMLabeler):
        self.file_lines = file_lines
        self.root_dir = os.path.abspath(root_dir or "")
        self.file_path = os.path.abspath(file_path or "")
        self.labeler = labeler

        self.nodes: dict[str, dict] = {}
        self.edges: list[tuple[str, str, str]] = []
        self._next_id = 1

        self.loop_stack: list[tuple[str, str]] = []  # (continue_target, break_target)
        self.switch_stack: list[str] = []  # break_target
        self.terminal_exits: list[str] = []

        # (node_id, deterministic_base_label, raw_block_text)
        self._block_requests: list[tuple[str, str, str]] = []

    def new_node(self, shape: str, label: str) -> str:
        nid = f"n{self._next_id}"
        self._next_id += 1
        self.nodes[nid] = {"shape": shape, "label": clean_label_text(label)}
        return nid

    def new_process_block(self, raw_block_text: str) -> str:
        nid = f"n{self._next_id}"
        self._next_id += 1

        raw = clean_unicode_chars(raw_block_text or "")[:MAX_BLOCK_CHARS_FOR_NODE]
        det = build_pseudocode_label(raw)
        self.nodes[nid] = {"shape": "process", "label": det}

        if self.labeler and self.labeler.enabled:
            # LLM will only lightly rewrite this deterministic label.
            self._block_requests.append((nid, det, raw))
        return nid

    def add_edge(self, src: str, dst: str, label: str = ""):
        if not src or not dst:
            return
        self.edges.append((src, dst, clean_label_text(label) if label else ""))

    def build_function(self, fn_cursor) -> Graph:
        body = None
        for ch in fn_cursor.get_children():
            if ch.kind == cindex.CursorKind.COMPOUND_STMT:
                body = ch
                break

        if body is None:
            n = self.new_node("process", "No implementation")
            self.add_edge("Start", n)
            self.add_edge(n, "End")
            return Graph(self.nodes, self.edges)

        entry, exits = self.build_compound(body)
        self.add_edge("Start", entry)
        for ex in exits:
            self.add_edge(ex, "End")
        for ex in self.terminal_exits:
            self.add_edge(ex, "End")

        # Batched label refinement (single call for many blocks)
        if self._block_requests and self.labeler and self.labeler.enabled:
            mapping = self.labeler.label_many(self._block_requests)
            for nid, lbl in mapping.items():
                if nid in self.nodes and self.nodes[nid].get("shape") == "process":
                    self.nodes[nid]["label"] = clean_label_text(lbl)

        return Graph(self.nodes, self.edges)

    def build_stmt(self, cursor) -> tuple[str, list[str]]:
        k = cursor.kind

        if _is_range_for_cursor(cursor):
            return self._build_range_for(cursor)

        if k == cindex.CursorKind.RETURN_STMT:
            return_text = cursor_text(cursor, self.file_lines).strip()
            # Handle `return;` (no whitespace) as well as `return expr;`
            return_expr = re.sub(r"^\s*return\b", "", return_text, flags=re.IGNORECASE).strip()
            return_expr = re.sub(r";+\s*$", "", return_expr).strip()
            label = f"return {return_expr}".strip() if return_expr else "return"
            n = self.new_node("process", label)
            self.terminal_exits.append(n)
            return n, []

        if k == cindex.CursorKind.BREAK_STMT:
            n = self.new_node("process", "break")
            if self.switch_stack:
                self.add_edge(n, self.switch_stack[-1])
            elif self.loop_stack:
                self.add_edge(n, self.loop_stack[-1][1])
            return n, []

        if k == cindex.CursorKind.CONTINUE_STMT:
            n = self.new_node("process", "continue")
            if self.loop_stack:
                self.add_edge(n, self.loop_stack[-1][0])
            return n, []

        if k == cindex.CursorKind.IF_STMT:
            return self._build_if(cursor)

        if k in (cindex.CursorKind.FOR_STMT, cindex.CursorKind.WHILE_STMT, cindex.CursorKind.DO_STMT):
            header_text = cursor_text(cursor, self.file_lines)
            if k == cindex.CursorKind.FOR_STMT and re.search(r"for\s*\([^)]*:[^)]*\)", header_text):
                return self._build_range_for(cursor)
            return self._build_loop(cursor)

        if CK_CXX_FOR_RANGE is not None and k == CK_CXX_FOR_RANGE:
            return self._build_range_for(cursor)

        if k == cindex.CursorKind.SWITCH_STMT:
            return self._build_switch(cursor)

        if CK_CXX_TRY is not None and k == CK_CXX_TRY:
            return self._build_try(cursor)

        return self.build_compound(cursor)

    def build_compound(self, cursor) -> tuple[str, list[str]]:
        if cursor is None:
            n = self.new_node("process", "No operation")
            return n, [n]

        if cursor.kind != cindex.CursorKind.COMPOUND_STMT:
            text = cursor_text(cursor, self.file_lines)
            n = self.new_process_block(text)
            return n, [n]

        children = list(cursor.get_children())
        entry = None
        curr_exits: list[str] = []
        pending: list = []

        CONTROL_KINDS = {
            cindex.CursorKind.IF_STMT,
            cindex.CursorKind.FOR_STMT,
            cindex.CursorKind.WHILE_STMT,
            cindex.CursorKind.DO_STMT,
            CK_CXX_FOR_RANGE,
            cindex.CursorKind.SWITCH_STMT,
            CK_CXX_TRY,
            cindex.CursorKind.RETURN_STMT,
            cindex.CursorKind.BREAK_STMT,
            cindex.CursorKind.CONTINUE_STMT,
        }
        CONTROL_KINDS = {x for x in CONTROL_KINDS if x is not None}

        def _is_control_stmt(c) -> bool:
            if c is None:
                return False
            if _is_range_for_cursor(c):
                return True
            if c.kind == cindex.CursorKind.FOR_STMT:
                ht = cursor_text(c, self.file_lines)
                if re.search(r"for\s*\([^)]*:[^)]*\)", ht):
                    return True
            return c.kind in CONTROL_KINDS

        def flush_pending():
            nonlocal entry, curr_exits, pending
            if not pending:
                return

            segments: list[list] = []
            current: list = []
            current_lines = 0

            for stmt in pending:
                stmt_text = cursor_text(stmt, self.file_lines)
                stmt_lines = max(1, len(stmt_text.splitlines())) if stmt_text else 1

                if current and (
                    len(current) >= PENDING_SEGMENT_MAX_STMTS
                    or (current_lines + stmt_lines) > PENDING_SEGMENT_MAX_LINES
                ):
                    segments.append(current)
                    current = []
                    current_lines = 0

                current.append(stmt)
                current_lines += stmt_lines

            if current:
                segments.append(current)

            for seg in segments:
                text = "\n".join(t for t in (cursor_text(s, self.file_lines) for s in seg) if t)
                n = self.new_process_block(text)

                if entry is None:
                    entry = n
                if curr_exits:
                    for ex in curr_exits:
                        self.add_edge(ex, n)
                curr_exits = [n]

            pending = []

        for child in children:
            if _is_control_stmt(child):
                flush_pending()
                s_entry, s_exits = self.build_stmt(child)

                if entry is None:
                    entry = s_entry
                if curr_exits:
                    for ex in curr_exits:
                        self.add_edge(ex, s_entry)
                curr_exits = s_exits
            else:
                pending.append(child)

        flush_pending()

        if entry is None:
            n = self.new_node("process", "No operation")
            return n, [n]

        return entry, curr_exits

    def _build_if(self, cursor) -> tuple[str, list[str]]:
        children = list(cursor.get_children())
        cond = children[0] if children else None
        then_stmt = children[1] if len(children) > 1 else None
        else_stmt = children[2] if len(children) > 2 else None

        cond_text = cursor_extent_text(cond, self.file_lines) or cursor_text(cond, self.file_lines) or "condition"
        cond_label = clean_condition_text(cond_text)
        d = self.new_node("decision", f"Check if {cond_label}" if cond_label else "Check condition")

        t_entry, t_exits = self.build_compound(then_stmt)
        self.add_edge(d, t_entry, "true")

        if else_stmt:
            f_entry, f_exits = self.build_compound(else_stmt)
            self.add_edge(d, f_entry, "false")
            return d, (t_exits + f_exits)

        return d, (t_exits + [d])

    def _build_loop(self, cursor) -> tuple[str, list[str]]:
        k = cursor.kind
        children = list(cursor.get_children())

        cond_text = ""
        cond = None
        body = None
        if k == cindex.CursorKind.WHILE_STMT:
            cond = children[0] if len(children) > 0 else None
            body = children[1] if len(children) > 1 else None
            cond_text = cursor_extent_text(cond, self.file_lines) or cursor_text(cond, self.file_lines)
        elif k == cindex.CursorKind.DO_STMT:
            body = children[0] if len(children) > 0 else None
            cond = children[1] if len(children) > 1 else None
            cond_text = cursor_extent_text(cond, self.file_lines) or cursor_text(cond, self.file_lines)
        else:  # FOR
            header_text = cursor_text(cursor, self.file_lines)
            inside = _extract_paren_group_after("for", header_text)
            if inside and ";" in inside:
                parts = [p.strip() for p in _split_top_level_semicolons(inside)]
                cond_text = parts[1] if len(parts) >= 2 else ""
                cond = None
            else:
                cond = children[1] if len(children) > 1 else None
                cond_text = cursor_extent_text(cond, self.file_lines) or cursor_text(cond, self.file_lines)
            body = children[-1] if children else None

        cond_text = (cond_text or cursor_text(cond, self.file_lines) or "loop condition").strip()
        cond_label = clean_condition_text(cond_text)

        # For-loop phrasing: "for all X less than Y" (no operators)
        if k == cindex.CursorKind.FOR_STMT and inside and ";" in inside:
            init_part = parts[0] if len(parts) > 0 else ""
            cond_part = parts[1] if len(parts) > 1 else cond_text
            vmatch = re.search(r"\b([A-Za-z_][A-Za-z0-9_]*)\s*=", init_part)
            var = vmatch.group(1) if vmatch else ""
            cm = re.search(r"\b([A-Za-z_][A-Za-z0-9_:\.\[\]]+)\s*(<|<=|>|>=)\s*([A-Za-z_][A-Za-z0-9_:\.\[\]]+)", cond_part)
            if var and cm:
                op = cm.group(2)
                left = cm.group(1).strip()
                right = cm.group(3).strip()
                # Normalize direction: prefer "var less than bound"
                bound = right
                if op in (">", ">="):
                    bound = left
                bound_txt = _describe_member_chain(bound) or bound
                vtxt = _humanize_var(var)
                # simple pluralization for "types"/"indexes"
                if not vtxt.endswith("s") and vtxt.lower().endswith("type"):
                    vtxt = vtxt[:-4] + "types"
                label_txt = f"for all {vtxt} less than {bound_txt}"
                check = self.new_node("decision", label_txt)
            else:
                check = self.new_node("decision", f"Check {cond_label}" if cond_label else "Check loop condition")
        else:
            # while/do: use "Check ..." phrasing
            check = self.new_node("decision", f"Check {cond_label}" if cond_label else "Check loop condition")
        after = self.new_node("process", "After loop")

        self.loop_stack.append((check, after))
        b_entry, b_exits = self.build_compound(body)
        self.loop_stack.pop()

        if k == cindex.CursorKind.DO_STMT:
            entry = b_entry
            for ex in b_exits:
                self.add_edge(ex, check)
            self.add_edge(check, b_entry, "true")
            self.add_edge(check, after, "false")
            return entry, [after]

        self.add_edge(check, b_entry, "true")
        self.add_edge(check, after, "false")
        for ex in b_exits:
            self.add_edge(ex, check)
        return check, [after]

    def _build_range_for(self, cursor) -> tuple[str, list[str]]:
        children = list(cursor.get_children())
        body = children[-1] if children else None

        header_text = cursor_text(cursor, self.file_lines)
        inside = _extract_paren_group_after("for", header_text)
        cond_text = inside if inside else "item : range"
        cond_label = clean_condition_text(cond_text)

        check = self.new_node("decision", f"for each {cond_label}" if cond_label else "for each item")
        after = self.new_node("process", "After loop")

        self.loop_stack.append((check, after))
        b_entry, b_exits = self.build_compound(body)
        self.loop_stack.pop()

        self.add_edge(check, b_entry, "true")
        self.add_edge(check, after, "false")
        for ex in b_exits:
            self.add_edge(ex, check)
        return check, [after]

    def _build_switch(self, cursor) -> tuple[str, list[str]]:
        children = list(cursor.get_children())
        expr = children[0] if children else None
        body = children[1] if len(children) > 1 else None

        expr_text = cursor_extent_text(expr, self.file_lines) or cursor_text(expr, self.file_lines) or "expression"
        expr_label = clean_condition_text(expr_text)
        d = self.new_node("decision", f"switch on {expr_label}" if expr_label else "switch")
        after = self.new_node("process", "After switch")

        self.switch_stack.append(after)

        cases = []
        default_case = None
        if body:
            for ch in body.get_children():
                if ch.kind == cindex.CursorKind.CASE_STMT:
                    cases.append(ch)
                elif ch.kind == cindex.CursorKind.DEFAULT_STMT:
                    default_case = ch

        built: list[tuple[str, str, list[str]]] = []
        for c in cases:
            raw = cursor_text(c, self.file_lines)
            case_label = clean_label_text(raw.split(":")[0] if ":" in raw else "case")
            c_children = list(c.get_children())
            c_body = c_children[-1] if c_children else None
            entry, exits = self.build_compound(c_body)
            built.append((case_label, entry, exits))

        if default_case:
            d_children = list(default_case.get_children())
            d_body = d_children[-1] if d_children else None
            entry, exits = self.build_compound(d_body)
            built.append(("default", entry, exits))

        for lbl, entry, _ in built:
            self.add_edge(d, entry, lbl)

        # Best-effort fallthrough
        for i in range(len(built) - 1):
            next_entry = built[i + 1][1]
            for ex in built[i][2]:
                self.add_edge(ex, next_entry, "fallthrough")

        if built:
            for ex in built[-1][2]:
                self.add_edge(ex, after)
        else:
            self.add_edge(d, after, "default")

        self.switch_stack.pop()
        return d, [after]

    def _build_try(self, cursor) -> tuple[str, list[str]]:
        children = list(cursor.get_children())
        try_block = children[0] if children else None
        catches = children[1:] if len(children) > 1 else []

        if CK_CXX_CATCH is not None:
            catches = [c for c in catches if c.kind == CK_CXX_CATCH] or catches

        decision = self.new_node("decision", "Exception occurs")
        after = self.new_node("process", "After try/catch")

        t_entry, t_exits = self.build_compound(try_block)
        self.add_edge(decision, t_entry, "no")
        for ex in t_exits:
            self.add_edge(ex, after)

        if catches:
            catch_text = "\n".join(t for t in (cursor_text(c, self.file_lines) for c in catches) if t)
            c_node = self.new_process_block(catch_text or "Handle exception")
            self.add_edge(decision, c_node, "yes")
            self.add_edge(c_node, after)
        else:
            self.add_edge(decision, after, "yes")

        return decision, [after]


def _strip_internal_label_prefixes(label: str) -> str:
    """
    Strip internal control markers that should never appear in Mermaid labels.
    Applied ONLY at Mermaid rendering time.
    """
    s = clean_unicode_chars(label or "").strip()
    if not s:
        return ""

    # Unwrap wrappers like: [[Process; ...]] / [[Decision: ...]]
    m = re.match(r"^\[\[\s*(process|decision)\s*[:;,\-]\s*(.*?)\s*\]\]\s*$", s, flags=re.IGNORECASE | re.DOTALL)
    if m:
        return (m.group(2) or "").strip()

    # Drop leading prefixes: [Process] / [Decision]
    s = re.sub(r"^\[\s*(process|decision)\s*\]\s*", "", s, flags=re.IGNORECASE)
    # Drop leading internal wrappers even if closing ']]' was lost upstream.
    s = re.sub(r"^\[\[\s*(process|decision)\s*[:;,\-]\s*", "", s, flags=re.IGNORECASE)
    # Drop trailing closing wrapper if present.
    s = re.sub(r"\s*\]\]\s*$", "", s)
    return s.strip()


def render_mermaid(graph: Graph) -> str:
    lines = ["flowchart TD", "Start((Start))"]

    for nid, node in graph.nodes.items():
        shape = node["shape"]
        label = clean_label_text(_strip_internal_label_prefixes(node["label"])).replace("<br/>", "\n")
        if shape == "decision":
            lines.append(f"{nid}{{{{{label}}}}}")
        else:
            lines.append(f"{nid}[{label}]")

    lines.append("End((End))")

    for src, dst, lbl in graph.edges:
        if lbl:
            lines.append(f"{src} --> |{clean_label_text(lbl)}| {dst}")
        else:
            lines.append(f"{src} --> {dst}")

    return "\n".join(lines) + "\n"


def validate_mermaid(mermaid: str) -> tuple[bool, Optional[str]]:
    if not mermaid or not mermaid.strip():
        return False, "Empty flowchart"
    if "flowchart" not in mermaid.lower():
        return False, "Missing flowchart declaration"
    if "Start((Start))" not in mermaid:
        return False, "Missing Start node"
    if "End((End))" not in mermaid:
        return False, "Missing End node"
    # Allow #40; codes, but do not allow HTML entities
    if "&#" in mermaid:
        return False, "HTML entities detected (use #40; style codes instead)"
    for line in mermaid.splitlines():
        if line.count("-->") > 1:
            return False, "Multiple edges in one line"
    return True, None


def generate_flowchart_for_function(fn_cursor, file_lines: list[str], root_dir: str, file_path: str):
    builder = FlowBuilder(file_lines, root_dir=root_dir, file_path=file_path, labeler=_GLOBAL_LABELER)
    graph = builder.build_function(fn_cursor)
    mermaid = render_mermaid(graph)

    ok, err = validate_mermaid(mermaid)
    if not ok:
        return (mermaid, None, f"Flowchart validation failed: {err}")

    # Optional image generation
    img = None
    if not _GLOBAL_MERMAID_DIR or not os.path.isdir(_GLOBAL_MERMAID_DIR):
        return (mermaid, None, "Mermaid image skipped: mermaid converter directory not found")

    currdir = os.getcwd()
    try:
        os.chdir(_GLOBAL_MERMAID_DIR)
        subprocess.check_output(
            ["node", "index.js", mermaid, f"{fn_cursor.spelling}.png"],
            stderr=subprocess.STDOUT,
            timeout=30,
        )
        img = os.path.join(_GLOBAL_MERMAID_DIR, f"{fn_cursor.spelling}.png")
        return (mermaid, img, None)
    except Exception as e:
        return (mermaid, None, f"Mermaid image generation failed: {str(e)[:200]}")
    finally:
        os.chdir(currdir)


def generate_function_description(function_lines: list[str]) -> str:
    if _GLOBAL_NO_DESC or not _GLOBAL_LLM or not HumanMessage:
        return ""
    prompt = (
        "You are a C++ code documentation expert.\n"
        "Provide a concise 2-3 sentence description.\n"
        "Do not invent anything.\n\n"
        "Function:\n"
        "{function}\n"
        "Description:"
    )
    query = prompt.format(function="\n".join(function_lines[:120]))
    try:
        log("[LLM] Description prompt:\n" + query)
        t0 = time.perf_counter()
        resp = _GLOBAL_LLM.invoke([HumanMessage(query)])
        dt = time.perf_counter() - t0
        log(f"[LLM] Description completed in {dt:.3f}s")
        return clean_unicode_chars(resp.content).strip()
    except Exception:
        return ""


def extract_node_info(fn_cursor, file_path: str, module_name: str, root_dir: str) -> Optional[dict]:
    extent = fn_cursor.extent
    try:
        fn_t0 = time.perf_counter()
        log(f"[INFO] Function: {fn_cursor.spelling} ({file_path}:{extent.start.line}-{extent.end.line})")
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            file_lines = f.readlines()

        function_lines = extract_source_lines(file_lines, extent.start.line, extent.end.line)
        function_lines = [l.rstrip() for l in function_lines if l.strip()]
        if not function_lines:
            return None

        flowchart, img, feedback = generate_flowchart_for_function(
            fn_cursor, file_lines, root_dir=root_dir, file_path=file_path
        )

        result = {
            "uid": node_uid(fn_cursor),
            "name": fn_cursor.spelling,
            "line_start": extent.start.line,
            "column_start": extent.start.column,
            "line_end": extent.end.line,
            "column_end": extent.end.column,
            "file_name": file_path,
            "module_name": module_name,
            "description": generate_function_description(function_lines),
            "flowchart": flowchart,
            "feedback": feedback,
            "img": img,
            "callees": [],
            "callers": [],
        }
        fn_dt = time.perf_counter() - fn_t0
        log(f"[TIME] Function {fn_cursor.spelling} processed in {fn_dt:.3f}s")
        return result
    except Exception as e:
        print(f"[WARN] extract_node_info failed for {fn_cursor.spelling}: {e}")
        return None


def visit(
    cursor,
    file_path: str,
    module_name: str,
    root_dir: str,
    nodes: dict,
    call_edges,
    current_fn_uid: Optional[str],
    visited=None,
):
    if visited is None:
        visited = set()

    if cursor.location.file and cursor.location.file.name != file_path:
        return

    fqn = f"{module_name}::{file_path}::{cursor.spelling}"
    if fqn in visited:
        return

    if cursor.is_definition() and cursor.kind in (
        cindex.CursorKind.FUNCTION_DECL,
        cindex.CursorKind.CXX_METHOD,
    ):
        visited.add(fqn)
        uid = node_uid(cursor)
        if uid not in nodes and cursor.spelling:
            info = extract_node_info(cursor, file_path, module_name, root_dir=root_dir)
            if info:
                nodes[uid] = info
                current_fn_uid = uid

    if cursor.kind == cindex.CursorKind.CALL_EXPR and current_fn_uid:
        ref = cursor.referenced
        if ref and ref.spelling and ref.location.file:
            callee_uid = node_uid(ref)
            call_edges[current_fn_uid].add(callee_uid)

    for child in cursor.get_children():
        visit(child, file_path, module_name, root_dir, nodes, call_edges, current_fn_uid, visited)


def generate_word_document(data: list[dict], doc_name: str):
    if not Document or not Inches:
        log("[WARN] python-docx not installed; skipping DOCX generation")
        return

    doc = Document()
    for index, item in enumerate(data, start=1):
        doc.add_heading(f"1.1.{index} {item['name']}", level=1)
        table = doc.add_table(rows=2, cols=2, style="Table Grid")
        table.rows[0].cells[0].text = "Requirement ID"
        table.rows[0].cells[1].text = f"SAVV8-SwU-{index}"
        table.rows[1].cells[0].text = "Flowchart"

        if item.get("img") and os.path.exists(item["img"]):
            table.rows[1].cells[1].add_paragraph().add_run().add_picture(item["img"], width=Inches(6.0))
        else:
            table.rows[1].cells[1].text = item.get("feedback") or "Flowchart image not available"

    os.makedirs(os.path.dirname(doc_name), exist_ok=True)
    doc.save(doc_name)


def parse_file(index, file_path: str, root_dir: str, compile_args: list[str], out_nodes: dict, out_edges):
    module_name = get_module_name(file_path, root_dir)
    log(f"[INFO] Parsing file: {file_path}")
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            file_lines = f.readlines()
    except Exception:
        file_lines = []
    tu = index.parse(
        file_path,
        args=compile_args,
        options=cindex.TranslationUnit.PARSE_DETAILED_PROCESSING_RECORD,
    )

    # Build function definition snippets for LLM-based "purpose" inference.
    # We do NOT generate callee flowcharts here; snippets are used only as context.
    def _walk_defs(cur):
        try:
            if cur is None:
                return
            if cur.kind in (cindex.CursorKind.FUNCTION_DECL, cindex.CursorKind.CXX_METHOD) and cur.is_definition():
                nm = cur.spelling or ""
                if nm and nm not in _GLOBAL_FN_DEF_SNIPPET:
                    snip = _function_def_snippet(cur, file_lines)
                    if snip:
                        _GLOBAL_FN_DEF_SNIPPET[nm] = snip
            for ch in cur.get_children():
                _walk_defs(ch)
        except Exception:
            return

    _walk_defs(tu.cursor)

    my_nodes: dict = {}
    my_edges = defaultdict(set)
    visit(tu.cursor, file_path, module_name, root_dir, my_nodes, my_edges, None)

    if not my_nodes:
        return

    base = os.path.splitext(os.path.basename(file_path))[0]
    meta_name = base
    if meta_name in out_nodes:
        meta_name = f"{meta_name}_{len(out_nodes)}"

    json_path = os.path.join(_GLOBAL_OUT_DIR, f"{meta_name}.json")
    docx_path = os.path.join(_GLOBAL_OUT_DIR, f"{base}.docx")

    os.makedirs(_GLOBAL_OUT_DIR, exist_ok=True)
    generate_word_document(list(my_nodes.values()), docx_path)
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(list(my_nodes.values()), f, indent=2, ensure_ascii=False)

    out_nodes.update(my_nodes)
    out_edges.update(my_edges)


def parse_codebase(path: str, compile_args: Optional[list[str]] = None) -> list[dict]:
    """
    Parse a C/C++ codebase path.
    - If `path` is a directory: walk and parse all supported files.
    - If `path` is a single file: parse just that file.
    """
    compile_args = compile_args or ["-std=c++17"]
    index = cindex.Index.create()
    nodes: dict = {}
    call_edges = defaultdict(set)

    if os.path.isfile(path):
        if is_cpp_file(path):
            root_dir = os.path.dirname(os.path.abspath(path))
            try:
                parse_file(index, os.path.abspath(path), root_dir, compile_args, nodes, call_edges)
            except Exception as e:
                print(f"[WARN] Failed to parse {path}: {e}")
        return list(nodes.values())

    root_dir = path
    for root, _, files in os.walk(root_dir):
        for f in files:
            if is_cpp_file(f):
                file_path = os.path.join(root, f)
                try:
                    parse_file(index, file_path, root_dir, compile_args, nodes, call_edges)
                except Exception as e:
                    print(f"[WARN] Failed to parse {file_path}: {e}")
    return list(nodes.values())


# Globals configured in main()
_GLOBAL_OUT_DIR = DEFAULT_OUT_DIR
_GLOBAL_MERMAID_DIR = DEFAULT_MERMAID_DIR
_GLOBAL_LLM = None
_GLOBAL_LABELER: BatchLLMLabeler = BatchLLMLabeler(None, enabled=False)
_GLOBAL_NO_DESC = False
_GLOBAL_FN_DEF_SNIPPET: dict[str, str] = {}


def main():
    parser = argparse.ArgumentParser(description="Generate AST-driven flowcharts for C++ functions")
    parser.add_argument("--path", required=True, help="C++ codebase root directory OR a single .cpp/.c file")
    parser.add_argument("--std", default="c++17", help="C++ standard, e.g. c++17, c++20")

    parser.add_argument("--out-dir", default=DEFAULT_OUT_DIR, help="Output directory for json/docx/images")
    parser.add_argument("--mermaid-dir", default=DEFAULT_MERMAID_DIR, help="Directory containing index.js renderer")
    parser.add_argument("--no-desc", action="store_true", help="Skip function description generation")
    parser.add_argument("--ollama-model", default="gpt-oss", help="Ollama model name (default: gpt-oss)")
    args = parser.parse_args()

    global VERBOSE, _GLOBAL_OUT_DIR, _GLOBAL_MERMAID_DIR, _GLOBAL_LLM, _GLOBAL_LABELER, _GLOBAL_NO_DESC
    VERBOSE = True
    _GLOBAL_OUT_DIR = os.path.abspath(args.out_dir)
    _GLOBAL_MERMAID_DIR = os.path.abspath(args.mermaid_dir) if args.mermaid_dir else ""
    _GLOBAL_NO_DESC = bool(args.no_desc)

    # Mandatory batched LLM labeler (per requirements)
    _GLOBAL_LLM = ChatOllama(model=args.ollama_model, temperature=0.1, top_k=10, top_p=0.9)
    _GLOBAL_LABELER = BatchLLMLabeler(_GLOBAL_LLM, enabled=True)

    os.makedirs(_GLOBAL_OUT_DIR, exist_ok=True)
    all_t0 = time.perf_counter()
    parse_codebase(args.path, compile_args=[f"-std={args.std}"])
    all_dt = time.perf_counter() - all_t0
    log(f"[TIME] Total execution time: {all_dt:.3f}s")
    print("Done.")


if __name__ == "__main__":
    main()

